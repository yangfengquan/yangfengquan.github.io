import CoolProp.CoolProp as CP
import io
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Material import MaterialManager
# 设置中文字体
plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False
class FluidPipeAnalyzer:
    """流体管道分析器 - 支持多种流体"""
    """
    Pressure: Pa
    temperature: K
    length: m
    Energy: j
    power: W
    """
    def __init__(self):
        self.results = []
        self.material_manager = MaterialManager()
        
    def get_fluid_properties(self, P_pa, T_k=None, enthalpy=None, quality=None, fluid='Water'):
        """获取流体物性 - 支持多种流体"""
        try:
            if quality is not None and 0 <= quality <= 1:
                T_k = CP.PropsSI('T', 'P', P_pa, 'Q', quality, fluid)
                h = CP.PropsSI('H', 'P', P_pa, 'Q', quality, fluid)
            elif enthalpy is not None:
                T_k = CP.PropsSI('T', 'P', P_pa, 'H', enthalpy, fluid)
                h = enthalpy
                quality = CP.PropsSI('Q', 'P', P_pa, 'H', enthalpy, fluid)
            else:
                h = CP.PropsSI('H', 'P', P_pa, 'T', T_k, fluid)
                quality = CP.PropsSI('Q', 'P', P_pa, 'T', T_k, fluid)
            
            # 获取其他物性
            rho = CP.PropsSI('D', 'P', P_pa, 'H', h, fluid)
            mu = CP.PropsSI('V', 'P', P_pa, 'H', h, fluid)
            k = CP.PropsSI('L', 'P', P_pa, 'H', h, fluid)
            cp = CP.PropsSI('C', 'P', P_pa, 'H', h, fluid)
            
            # 对于非水流体，尝试获取饱和温度
            try:
                T_sat = CP.PropsSI('T', 'P', P_pa, 'Q', 0, fluid) 
            except:
                T_sat = 100.0  # 默认值
            
            # 相态判断
            if quality <= 0:
                phase = "subcooled_liquid" if T_k < T_sat else "saturated_liquid"
            elif quality >= 1:
                phase = "superheated_vapor" if T_k > T_sat else "saturated_vapor"
            else:
                phase = "two_phase"
            
            return {
                'pressure_pa': P_pa,
                'temperature_k': T_k,
                'density_kg_per_m3': rho,
                'viscosity_Pa_s': mu,
                'enthalpy_J_per_kg': h,
                'conductivity_W_per_mK': k,
                'specific_heat_J_per_kgK': cp,
                'quality': max(0.0, min(1.0, quality)),
                'saturation_temperature_c': T_sat,
                'phase': phase,
                'superheat_k': max(0, T_k - T_sat),
                'subcooling_k': max(0, T_sat - T_k)
            }
            
        except Exception as e:
            print(f"流体物性计算失败 ({fluid}): {e}")
            
    def friction_factor(self, Re, roughness, diameter):
        """摩擦系数计算 - Swamee-Jain公式"""
        if Re <= 0:
            return 0.015
        
        if Re < 2000:
            return 64.0 / Re
        
        if Re < 4000:
            f_lam = 64.0 / 2000
            f_turb = self._swamee_jain(4000, roughness, diameter)
            weight = (Re - 2000) / 2000
            return (1 - weight) * f_lam + weight * f_turb
        
        return self._swamee_jain(Re, roughness, diameter)

    def _swamee_jain(self, Re, roughness, diameter):
        """Swamee-Jain公式"""
        eD = roughness / diameter
        f = 0.25 / (np.log10(eD/3.7 + 5.74/Re**0.9))**2
        return f

    def pressure_drop_calculation(self, props, mass_flow, diameter, length, 
                                roughness, fittings_resistance=0):
        """压力降计算 - 包含局部阻力"""
        area = np.pi * (diameter/2)**2
        rho = props['density_kg_per_m3']
        mu = props['viscosity_Pa_s']

        velocity = mass_flow / (rho * area)
        Re = (rho * velocity * diameter) / mu
        f = self.friction_factor(Re, roughness, diameter)
        
        # 沿程阻力损失
        dp_friction = f * (length / diameter) * (rho * velocity**2) / 2
        
        # 局部阻力损失
        dp_fittings = fittings_resistance * (rho * velocity**2) / 2
        
        dp_total = dp_friction + dp_fittings

        return {
            'velocity_m_s': velocity,
            'reynolds_number': Re,
            'friction_factor': f,
            'pressure_drop_pa': dp_total,
            'friction_drop_pa': dp_friction,
            'fittings_drop_pa': dp_fittings
        }

    def heat_loss_calculation(self, fluid_temp_k, ambient_temp_k, pipe_od, 
                            insulation_thickness, insulation_material, 
                            protection_material, length,
                            wind_speed):
        """热损失计算 - 包含外保护层"""

        if insulation_thickness <= 0:
            # 无保温情况
            h_combined = self._calculate_external_heat_transfer(
                ambient_temp_k, ambient_temp_k, wind_speed, protection_material.emissivity
            )
            surface_area = np.pi * pipe_od * length
            Q_total = h_combined * surface_area * (fluid_temp_k - ambient_temp_k)
            Q_per_m = Q_total / length
            surface_temp_k = ambient_temp_k + 0.7 * (fluid_temp_k - ambient_temp_k)
            # 计算单位外表面积热损失
            surface_area_per_m = np.pi * d_outer
            Q_per_area = Q_per_m / surface_area_per_m if surface_area_per_m > 0 else 0
            return {
                'total_heat_loss_w': Q_total,
                'heat_loss_per_m_w': Q_per_m,
                'heat_loss_per_area_w': Q_per_area,
                'overall_heat_transfer_coeff': h_combined,
                'convection_coeff': h_combined * 0.8,
                'surface_temp_k': surface_temp_k
            }
        
        # 有保温情况 - 迭代计算
        d_outer = pipe_od + 2 * insulation_thickness
        
        # 初始假设表面温度
        surface_temp_k = ambient_temp_k + 0.3 * (fluid_temp_k - ambient_temp_k)
        
        for iteration in range(20):
            # 计算保温层平均温度
            tm_k = (fluid_temp_k + surface_temp_k) / 2
            
            # 计算保温材料导热系数
            insulation_k = insulation_material.calculate_conductivity(tm_k)
            
            # 计算保温层热阻
            R_insulation = np.log(d_outer / pipe_od) / (2 * np.pi * insulation_k)
            
            # 计算外部换热系数
            h_external = self._calculate_external_heat_transfer(
                surface_temp_k, ambient_temp_k, wind_speed, protection_material.emissivity, pipe_od, d_outer
            )
            
            # 计算外部热阻
            R_external = 1 / (h_external * np.pi * d_outer)
            
            # 总热阻
            R_total = R_insulation + R_external
            
            # 单位长度热损失
            Q_per_m = (fluid_temp_k - ambient_temp_k) / R_total

            #总热损失
            Q_total = Q_per_m * length
            
            # 更新表面温度
            surface_temp_new_k = ambient_temp_k + Q_per_m * R_external
            
            # 收敛检查
            if abs(surface_temp_new_k - surface_temp_k) < 0.1:
                surface_temp_k = surface_temp_new_k
                break
                
            surface_temp_k = surface_temp_new_k
        
        # 计算对流换热系数
        h_convection = self._calculate_convection_coeff(wind_speed=wind_speed, surface_temp_k=surface_temp_k,
                                                        ambient_temp_k=ambient_temp_k, pipe_od=pipe_od, d_outer=d_outer)
        # 计算单位外表面积热损失
        surface_area_per_m = np.pi * d_outer
        Q_per_area = Q_per_m / surface_area_per_m if surface_area_per_m > 0 else 0
    
        return {
            'total_heat_loss_w': Q_total,
            'heat_loss_per_m_w': Q_per_m,
            'heat_loss_per_area_w': Q_per_area,
            'overall_heat_transfer_coeff': h_external,
            'convection_coeff': h_convection,
            'surface_temp_k': surface_temp_k,
            'insulation_conductivity': insulation_k
        }

    def _calculate_external_heat_transfer(self, surface_temp_k, ambient_temp_k, wind_speed, emissivity, pipe_od, d_outer):
        """计算外部总换热系数"""
        # 对流换热系数
        h_conv = self._calculate_convection_coeff(wind_speed=wind_speed, surface_temp_k=surface_temp_k,
                                                  ambient_temp_k=ambient_temp_k, pipe_od=pipe_od, d_outer=d_outer)
        
        # 辐射换热系数
        h_rad = self._calculate_radiation_coeff(surface_temp_k, ambient_temp_k, emissivity)
        
        return h_conv + h_rad

    def _calculate_convection_coeff(self, wind_speed, surface_temp_k=None, ambient_temp_k=None, pipe_od=None, d_outer=None):
        """计算对流换热系数"""
       
        if wind_speed == 0:
            #无风时
            h_conv = 26.4 / np.sqrt(297 - 273.15 + 0.5 * (surface_temp_k + ambient_temp_k)) * ((surface_temp_k - ambient_temp_k) / pipe_od)**0.25
        elif wind_speed * d_outer <= 0.8:
            # 自然对流
            h_conv = 0.08 / d_outer + 4.2 * wind_speed**0.618 / d_outer**0.382
        else:
            # 强制对流
            h_conv = 4.53 * wind_speed**0.805 / pipe_od**0.195
        
        return h_conv

    def _calculate_radiation_coeff(self, surface_temp_k, ambient_temp_k, emissivity):
        """计算辐射换热系数"""
        h_rad = 5.669 * emissivity / (surface_temp_k - ambient_temp_k) * ((surface_temp_k / 100)**4 - (ambient_temp_k / 100)**4) 
        
        return h_rad

    def analyze_pipe_segment(self, inlet_props, mass_flow, pipe_od, wall_thickness, length, 
                           pipe_type, insulation_thickness, insulation_material,
                           protection_material, ambient_temp_k, fittings_resistance=0,
                           wind_speed=3.0, fluid='Water'):
        """分析单个管段"""
        try:
            # 1. 压力降计算
            dp_result = self.pressure_drop_calculation(
                inlet_props, mass_flow, pipe_od-2*wall_thickness, length, pipe_type.roughness, fittings_resistance
            )
            
            outlet_pressure = max(10.0, inlet_props['pressure_pa'] - dp_result['pressure_drop_pa'])

            # 2. 热损失计算
            heat_loss_result = self.heat_loss_calculation(
                inlet_props['temperature_k'], ambient_temp_k, pipe_od,
                insulation_thickness, insulation_material, protection_material, length, wind_speed
            )
            
            # 3. 能量平衡
            delta_h = heat_loss_result['total_heat_loss_w'] / mass_flow if mass_flow > 0 else 0
            outlet_enthalpy = inlet_props['enthalpy_J_per_kg'] - delta_h
            
            # 4. 出口状态
            outlet_props = self.get_fluid_properties(outlet_pressure, enthalpy=outlet_enthalpy, fluid=fluid)
            
            # 5. 计算气相和液相流量
            vapor_flow = mass_flow * outlet_props['quality']
            liquid_flow = mass_flow * (1 - outlet_props['quality'])
            
            return outlet_props, dp_result, heat_loss_result, vapor_flow, liquid_flow
            
        except Exception as e:
            print(f"管段分析错误: {e}")
            return False

    def analyze(self, params):
        """分析"""
        pipe_name = params['pipe_name']
        inlet_pressure_pa = params['inlet_pressure_pa']
        inlet_temperature_k = params['inlet_temperature_k']
        inlet_quality = params.get('inlet_quality')
        mass_flow_kg_s = params['mass_flow_kg_s']
        pipe_length_m = params['pipe_length_m']
        pipe_od = params['pipe_od_m']
        pipe_wall_thickness_m = params['pipe_wall_thickness_m']
        pipe_type_name = params['pipe_type']
        insulation_thickness_m = params['insulation_thickness_m']
        insulation_material_name = params['insulation_material']
        protection_material_name = params['protection_material']
        ambient_temperature_k = params['ambient_temperature_k']
        wind_speed_m_s = params['wind_speed_m_s']
        segment_length_m = params['segment_length_m']
        fluid = params['fluid']
        fittings_data = params.get('fittings_data', [])
              
        # 获取材料
        pipe_type = self.material_manager.pipe_types.get(pipe_type_name)
        insulation_material = self.material_manager.insulation_materials.get(insulation_material_name)
        protection_material = self.material_manager.protection_materials.get(protection_material_name)
        if not insulation_material:
            print(f"警告: 未找到管道类别 '{insulation_material_name}'，使用默认值")
            pipe_type = list(self.material_manager.pipe_types.values())[2]
        
        if not insulation_material:
            print(f"警告: 未找到保温材料 '{insulation_material_name}'，使用默认值")
            insulation_material = list(self.material_manager.insulation_materials.values())[0]
        
        if not protection_material:
            print(f"警告: 未找到外保护层材料 '{protection_material_name}'，使用默认值")
            protection_material = list(self.material_manager.protection_materials.values())[0]
        
        # 计算总局部阻力系数
        total_fittings_resistance = 0
        for fitting_name, count in fittings_data:
            fitting = self.material_manager.pipe_fittings.get(fitting_name)
            if fitting:
                total_fittings_resistance += fitting.resistance_coef * count

        # 获取入口条件
        if inlet_quality is not None:
            inlet_props = self.get_fluid_properties(inlet_pressure_pa, quality=inlet_quality, fluid=fluid)
        else:
            inlet_props = self.get_fluid_properties(inlet_pressure_pa, T_k=inlet_temperature_k, fluid=fluid)
        
        # 计算入口流速
        area = np.pi * ((pipe_od - 2 * pipe_wall_thickness_m) / 2)**2
        inlet_velocity = mass_flow_kg_s / (inlet_props['density_kg_per_m3'] * area)

        current_props = inlet_props
        results = []
        
        num_segments = max(1, int(np.ceil(pipe_length_m / segment_length_m)))
        
        for i in range(num_segments):
            seg_length = min(segment_length_m, pipe_length_m - i * segment_length_m)
            distance = (i + 1) * seg_length
            
            current_props, dp_result, heat_loss_result, vapor_flow, liquid_flow = self.analyze_pipe_segment(
                current_props, mass_flow_kg_s, pipe_od, pipe_wall_thickness_m, seg_length,
                pipe_type, insulation_thickness_m, insulation_material, 
                protection_material, ambient_temperature_k, total_fittings_resistance / num_segments,
                wind_speed_m_s, fluid
            )
            
            result = {
                'segment': i+1,
                'distance_m': distance,
                'pressure_pa': current_props['pressure_pa'],
                'temperature_k': current_props['temperature_k'],
                'quality': current_props['quality'],
                'enthalpy_J_kg': current_props['enthalpy_J_per_kg'],
                'density_kg_m3': current_props['density_kg_per_m3'],
                'velocity_m_s': dp_result['velocity_m_s'],
                'pressure_drop_pa': dp_result['pressure_drop_pa'],
                'friction_drop_pa': dp_result['friction_drop_pa'],
                'fittings_drop_pa': dp_result['fittings_drop_pa'],
                'total_heat_loss_w': heat_loss_result['total_heat_loss_w'],
                'heat_loss_per_m_w': heat_loss_result['heat_loss_per_m_w'],
                'heat_loss_per_area_w': heat_loss_result['heat_loss_per_area_w'],
                'overall_heat_transfer_coeff': heat_loss_result['overall_heat_transfer_coeff'],#表面换热系数
                'convection_coeff': heat_loss_result['convection_coeff'],
                'surface_temp_k': heat_loss_result['surface_temp_k'],
                'insulation_conductivity': heat_loss_result['insulation_conductivity'],
                'reynolds_number': dp_result['reynolds_number'],
                'friction_factor': dp_result['friction_factor'],
                'vapor_flow_kg_s': vapor_flow,
                'liquid_flow_kg_s': liquid_flow
            }
            
            results.append(result)
            
        
        df_results = pd.DataFrame(results)
        
        # 完整结果分析
        analysis_results = self._comprehensive_results_analysis(
            df_results, inlet_props, pipe_length_m, mass_flow_kg_s, pipe_name, fluid,
            total_fittings_resistance, inlet_velocity, pipe_type.roughness
        )
        
        # 绘制完整图表
        fig = self._plot_comprehensive_results(df_results, pipe_name)
        
        return df_results, analysis_results, fig

    def _comprehensive_results_analysis(self, df, inlet_props, total_length_m, mass_flow, pipe_name, fluid, fittings_resistance, inlet_velocity_m_s, roughness_m):
        """完整结果分析"""
        final_pressure_pa = df['pressure_pa'].iloc[-1]
        final_temperature_k = df['temperature_k'].iloc[-1]
        final_quality = df['quality'].iloc[-1]
        final_velocity = df['velocity_m_s'].iloc[-1]
        final_vapor_flow_kg_s = df['vapor_flow_kg_s'].iloc[-1]
        final_liquid_flow_kg_s = df['liquid_flow_kg_s'].iloc[-1]
        
        total_pressure_drop_pa = inlet_props['pressure_pa'] - final_pressure_pa
        total_temperature_drop = inlet_props['temperature_k'] - final_temperature_k
        total_heat_loss_w = df['total_heat_loss_w'].sum()
        total_friction_drop_pa = df['friction_drop_pa'].sum()
        total_fittings_drop_pa = df['fittings_drop_pa'].sum()
        
        pressure_drop_pa_m = total_pressure_drop_pa / total_length_m
        pressure_ratio = total_pressure_drop_pa / inlet_props['pressure_pa']
        
        # 热损失参数
        avg_heat_loss_per_m = df['heat_loss_per_m_w'].mean()
        avg_heat_loss_per_area = df['heat_loss_per_area_w'].mean()
        max_surface_temp_k = df['surface_temp_k'].max()
        avg_surface_temp_k = df['surface_temp_k'].mean()
        
        # 流动参数
        max_velocity_m_s = df['velocity_m_s'].max()
        min_velocity_m_s = df['velocity_m_s'].min()
        avg_velocity_m_s = df['velocity_m_s'].mean()
        avg_reynolds = df['reynolds_number'].mean()
        avg_friction = df['friction_factor'].mean()

        analysis_results = {
            'pipe_name': pipe_name,
            'fluid': fluid,
            'inlet_pressure_pa': inlet_props['pressure_pa'],
            'inlet_temperature_k': inlet_props['temperature_k'],
            'inlet_quality': inlet_props['quality'],
            'inlet_velocity_m_s': inlet_velocity_m_s,
            'outlet_pressure_pa': final_pressure_pa,
            'outlet_temperature_k': final_temperature_k,
            'outlet_quality': final_quality,
            'outlet_velocity_m_s': final_velocity,
            'outlet_vapor_flow_kg_s': final_vapor_flow_kg_s,
            'outlet_liquid_flow_kg_s': final_liquid_flow_kg_s,
            'total_pressure_drop_pa': total_pressure_drop_pa,
            'pressure_drop_pa_m': pressure_drop_pa_m,
            'pressure_ratio': pressure_ratio,
            'total_temperature_drop_k': total_temperature_drop,
            'total_heat_loss_w': total_heat_loss_w,
            'total_friction_drop_pa': total_friction_drop_pa,
            'total_fittings_drop_pa': total_fittings_drop_pa,
            'fittings_resistance': fittings_resistance,
            'pipe_roughness_m': roughness_m,
            'avg_heat_loss_per_m_w': avg_heat_loss_per_m,
            'avg_heat_loss_per_area_w': avg_heat_loss_per_area,
            'max_surface_temp_k': max_surface_temp_k,
            'avg_surface_temp_k': avg_surface_temp_k,
            'max_velocity_m_s': max_velocity_m_s,
            'min_velocity_m_s': min_velocity_m_s,
            'avg_velocity_m_s': avg_velocity_m_s,
            'avg_reynolds': avg_reynolds,
            'avg_friction': avg_friction
        }
        
        return analysis_results
    
    def _plot_comprehensive_results(self, df, pipe_name):
        """绘制图表"""
        try:
            fig, axes = plt.subplots(3, 2, figsize=(15, 15))
            fig.suptitle(f'管道分析结果 - {pipe_name}', fontsize=16, fontweight='bold')
            
            # 使用彩色图表
            colors = ['blue', 'red', 'green', 'orange', 'purple', 'brown']
            line_styles = ['-', '--', '-.', ':', '-']
            markers = ['o', 's', '^', 'd', 'v']
            
            # 压力分布
            axes[0,0].plot(df['distance_m'], df['pressure_pa'] / 1000, 
                          color=colors[0], linestyle=line_styles[0], linewidth=2, 
                          marker=markers[0], markersize=4)
            axes[0,0].set_xlabel('距离 (m)')
            axes[0,0].set_ylabel('压力 (kPa)')
            axes[0,0].set_title('压力分布')
            axes[0,0].grid(True, alpha=0.3)
            
            # 温度和表面温度分布
            axes[0,1].plot(df['distance_m'], df['temperature_k'] - 273.15, 
                          color=colors[1], linestyle=line_styles[0], linewidth=2, 
                          marker=markers[1], markersize=4, label='流体温度')
            axes[0,1].plot(df['distance_m'], df['surface_temp_k'] - 273.15, 
                          color=colors[2], linestyle=line_styles[1], linewidth=2, 
                          label='表面温度')
            axes[0,1].set_xlabel('距离 (m)')
            axes[0,1].set_ylabel('温度 (°C)')
            axes[0,1].set_title('温度分布')
            axes[0,1].grid(True, alpha=0.3)
            axes[0,1].legend()
            
            # 干度分布
            axes[1,0].plot(df['distance_m'], df['quality'], 
                          color=colors[3], linestyle=line_styles[0], linewidth=2, 
                          marker=markers[2], markersize=4)
            axes[1,0].set_xlabel('距离 (m)')
            axes[1,0].set_ylabel('干度')
            axes[1,0].set_title('干度分布')
            axes[1,0].grid(True, alpha=0.3)
            axes[1,0].set_ylim(-0.05, 1.05)
            
            # 流速分布
            axes[1,1].plot(df['distance_m'], df['velocity_m_s'], 
                          color=colors[4], linestyle=line_styles[0], linewidth=2, 
                          marker=markers[3], markersize=4)
            axes[1,1].set_xlabel('距离 (m)')
            axes[1,1].set_ylabel('流速 (m/s)')
            axes[1,1].set_title('流速分布')
            axes[1,1].grid(True, alpha=0.3)
            
            # 单位长度热损失分布
            axes[2,0].plot(df['distance_m'], df['heat_loss_per_m_w'], 
                          color=colors[5], linestyle=line_styles[0], linewidth=2, 
                          marker=markers[4], markersize=4)
            axes[2,0].set_xlabel('距离 (m)')
            axes[2,0].set_ylabel('单位长度热损失 (W/m)')
            axes[2,0].set_title('单位长度热损失分布')
            axes[2,0].grid(True, alpha=0.3)
            
            # 累计热损失分布
            cumulative_heat_loss = df['total_heat_loss_w'].cumsum()
            axes[2,1].plot(df['distance_m'], cumulative_heat_loss, 
                          color=colors[0], linestyle=line_styles[0], linewidth=2)
            axes[2,1].set_xlabel('距离 (m)')
            axes[2,1].set_ylabel('累计热损失 (W)')
            axes[2,1].set_title('累计热损失分布')
            axes[2,1].grid(True, alpha=0.3)
            
            plt.tight_layout()
            return fig
            
        except Exception as e:
            print(f"绘图错误: {e}")
            return None
    
    def generate_word_report(self, params, df_results, analysis_results, fig, filename):
        """生成Word报告"""
        try:
            doc = Document()
            
            # 设置文档字体
            try:
                # 设置标题字体为黑体不加粗
                title_style = doc.styles['Title']
                title_style.font.name = '黑体'
                title_style.font.size = docx.shared.Pt(16)
                title_style.font.bold = False
                
                # 设置正文字体
                normal_style = doc.styles['Normal']
                normal_style.font.name = '宋体'
                normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                normal_style.font.size = docx.shared.Pt(10.5)
            except:
                pass
            
            # 标题
            title = doc.add_heading('管道系统分析报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 基本信息表格
            doc.add_heading('一、项目基本信息', level=1)
            basic_table = doc.add_table(rows=6, cols=2)
            basic_table.style = 'Table Grid'
            basic_table.autofit = False
            
            # 设置表格列宽
            for row in basic_table.rows:
                row.cells[0].width = Inches(2)
                row.cells[1].width = Inches(4)
            
            basic_data = [
                ["项目名称", params['pipe_name']],
                ["分析日期", pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')],
                ["流体介质", params['fluid']],
                ["管道总长度", f"{params['pipe_length_m']} m"],
                ["管道外径", f"{params['pipe_od_m']*1000} mm"],
                ["管道壁厚", f"{params['pipe_wall_thickness_m']*1000} mm"]
                
            ]
            
            for i, (key, value) in enumerate(basic_data):
                basic_table.cell(i, 0).text = key
                basic_table.cell(i, 1).text = str(value)
                # 设置单元格对齐方式
                basic_table.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                basic_table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            doc.add_paragraph()
            
            # 输入参数表格
            doc.add_heading('二、输入参数', level=1)
            input_table = doc.add_table(rows=8, cols=2)
            input_table.style = 'Table Grid'
            
            for row in input_table.rows:
                row.cells[0].width = Inches(2)
                row.cells[1].width = Inches(4)
            
            input_data = [
                ["入口压力", f"{params['inlet_pressure_pa'] / 1000} kPa (绝对压力)"],
                ["入口温度", f"{params['inlet_temperature_k']-273.15} °C"],
                ["质量流量", f"{params['mass_flow_kg_s']*3600:.1f} kg/h"],
                ["保温材料", params['insulation_material']],
                ["保温厚度", f"{params['insulation_thickness_m']*1000} mm"],
                ["外保护层", params['protection_material']],
                ["环境温度", f"{params['ambient_temperature_k']-273.15} °C"],
                ["风速", f"{params['wind_speed_m_s']} m/s"]
            ]
            
            for i, (key, value) in enumerate(input_data):
                input_table.cell(i, 0).text = key
                input_table.cell(i, 1).text = value
                input_table.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                input_table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            doc.add_paragraph()
            
            # 计算结果汇总表格
            doc.add_heading('三、计算结果汇总', level=1)
            result_table = doc.add_table(rows=11, cols=2)
            result_table.style = 'Table Grid'
            
            for row in result_table.rows:
                row.cells[0].width = Inches(2.5)
                row.cells[1].width = Inches(3.5)
            
            result_data = [
                ["出口压力", f"{analysis_results['outlet_pressure_pa']/1000:.1f} kPa"],
                ["出口温度", f"{analysis_results['outlet_temperature_k']-273.15:.1f} °C"],
                ["出口干度", f"{analysis_results['outlet_quality']:.4f}"],
                ["总压力降", f"{analysis_results['total_pressure_drop_pa']/1000:.1f} kPa"],
                ["单位压降", f"{analysis_results['pressure_drop_pa_m']:.1f} kPa/km"],
                ["压降比例", f"{analysis_results['pressure_ratio']:.3%}"],
                ["总热损失", f"{analysis_results['total_heat_loss_w']/1000:.1f} kW"],
                ["平均单位热损失", f"{analysis_results['avg_heat_loss_per_m_w']:.1f} W/m"],
                ["最大表面温度", f"{analysis_results['max_surface_temp_k']-273.15:.1f} °C"],
                ["入口流速", f"{analysis_results['inlet_velocity_m_s']:.2f} m/s"],
                ["出口流速", f"{analysis_results['outlet_velocity_m_s']:.2f} m/s"]
            ]
            
            for i, (key, value) in enumerate(result_data):
                result_table.cell(i, 0).text = key
                result_table.cell(i, 1).text = value
                result_table.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                result_table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            doc.add_paragraph()
            
            # 添加图表 - 单列显示
            if fig:
                doc.add_heading('四、分析图表', level=1)
                
                # 将图表保存到内存
                img_buffer = io.BytesIO()
                fig.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
                img_buffer.seek(0)
                
                # 添加到文档 - 单列显示，适当减小宽度
                doc.add_picture(img_buffer, width=Inches(6))
                doc.add_paragraph("图1 管道分析结果图表").alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 详细分段数据表
            doc.add_heading('五、详细分段计算结果', level=1)
            
            # 创建表格 - 调整列宽避免过宽
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Table Grid'
            table.autofit = False
            
            # 设置列宽 - 减小列宽避免表格过宽
            col_widths = [0.4, 0.6, 0.7, 0.7, 0.5, 0.7, 0.8, 0.8]  # 英寸
            for i, width in enumerate(col_widths):
                table.columns[i].width = Inches(width)
            
            # 表头 - 减少列数
            headers = ['段', '距离(m)', '压力(kPa)', '温度(°C)', '干度', '流速(m/s)', '热损失(W/m)', '表面温度(°C)']
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加数据（每10行显示一次）
            for i, row in df_results.iterrows():
                #if i % 10 == 0 or i == len(df_results) - 1:
                row_cells = table.add_row().cells
                row_data = [
                    f"{row['segment']:.0f}",
                    f"{row['distance_m']:.0f}",
                    f"{row['pressure_pa']/1000:.1f}",
                    f"{row['temperature_k']-273.15:.1f}",
                    f"{row['quality']:.4f}",
                    f"{row['velocity_m_s']:.2f}",
                    f"{row['heat_loss_per_m_w']:.1f}",
                    f"{row['surface_temp_k']-273.15:.1f}"
                ]
                
                for j, data in enumerate(row_data):
                    row_cells[j].text = data
                    row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # 技术分析
            doc.add_heading('六、技术分析', level=1)
            
            analysis_text = [
                "1. 压力损失分析:",
                f"   总压力损失为 {analysis_results['total_pressure_drop_pa']/1000:.1f} kPa，",
                f"   其中沿程阻力损失占 {analysis_results['total_friction_drop_pa']/analysis_results['total_pressure_drop_pa']:.1%}，",
                f"   局部阻力损失占 {analysis_results['total_fittings_drop_pa']/analysis_results['total_pressure_drop_pa']:.1%}。",
                "",
                "2. 热损失分析:",
                f"   管道总热损失为 {analysis_results['total_heat_loss_w']/1000:.1f} kW，",
                f"   平均单位长度热损失为 {analysis_results['avg_heat_loss_per_m_w']:.1f} W/m，",
                f"   平均单位外表面积热损失为 {analysis_results['avg_heat_loss_per_area_w']:.1f} W/m²。",
                "",
                "3. 流动特性分析:",
                f"   入口流速为 {analysis_results['inlet_velocity_m_s']:.2f} m/s，",
                f"   出口流速为 {analysis_results['outlet_velocity_m_s']:.2f} m/s，",
                f"   流动状态为 {'湍流' if analysis_results['avg_reynolds'] > 4000 else '层流'}，",
                f"   平均摩擦系数为 {analysis_results['avg_friction']:.5f}。",
                "",
                "4. 温度特性分析:",
                f"   最大外表面温度为 {analysis_results['max_surface_temp_k']-273.15:.1f} °C，",
                f"   平均外表面温度为 {analysis_results['avg_surface_temp_k']-273.15:.1f} °C。",
                "",
                "5. 相变分析:",
                f"   入口干度为 {analysis_results['inlet_quality']:.4f}，",
                f"   出口干度为 {analysis_results['outlet_quality']:.4f}，",
                f"   {'发生气化' if analysis_results['outlet_quality'] > analysis_results['inlet_quality'] else '发生凝结' if analysis_results['outlet_quality'] < analysis_results['inlet_quality'] else '未发生明显相变'}。"
            ]
            
            for text in analysis_text:
                if text.strip():
                    p = doc.add_paragraph(text)
                    if text.startswith("1.") or text.startswith("2.") or text.startswith("3.") or text.startswith("4.") or text.startswith("5."):
                        p.style = 'List Number'
            
            # 结论与建议
            doc.add_heading('七、结论与建议', level=1)
            
            conclusion_text = [
                "1. 系统性能评价:",
                f"   管道系统压降比例为 {analysis_results['pressure_ratio']:.2%}，",
                f"   {'在可接受范围内' if analysis_results['pressure_ratio'] < 0.05 else '偏高，建议优化管道设计'}。",
                "",
                "2. 保温效果评价:",
                f"   单位长度热损失为 {analysis_results['avg_heat_loss_per_m_w']:.1f} W/m，",
                f"   {'保温效果良好' if analysis_results['avg_heat_loss_per_m_w'] < 50 else '热损失较大，建议增加保温厚度'}。",
                "",
                "3. 安全性能评价:",
                f"   最大表面温度为 {analysis_results['max_surface_temp_k'] - 273.15:.1f} °C，",
                f"   {'符合安全要求' if analysis_results['max_surface_temp_k'] - 273.15 < 60 else '表面温度较高，需采取防护措施'}。"
            ]
            
            for text in conclusion_text:
                if text.strip():
                    p = doc.add_paragraph(text)
                    if text.startswith("1.") or text.startswith("2.") or text.startswith("3."):
                        p.style = 'List Number'
            
            # 保存文档
            doc.save(filename)
            return True
            
        except Exception as e:
            print(f"生成Word报告失败: {e}")
            return False
        